"""Generate the Tales Consulting discovery questionnaire for Fidson Healthcare as a .docx."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2D, 0x5C)
TEAL = RGBColor(0x2D, 0xCB, 0xC4)
GREY = RGBColor(0x55, 0x5F, 0x6E)
LIGHT_GREY = RGBColor(0xE5, 0xE7, 0xEB)

LOGO_PATH = Path(r"C:\Users\user\Desktop\tales consulting artifacts\tales_logo.png")
OUTPUT_PATH = Path(
    r"C:\Users\user\Desktop\tales consulting artifacts\Fidson_Discovery_Questionnaire.docx"
)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "BFC4CC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_cell_height(cell, cm: float) -> None:
    tr = cell._tc.getparent()
    tr_pr = tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), str(int(cm * 567)))
    height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(height)


def add_styled_run(paragraph, text: str, *, bold=False, size=11, color=None, font="Calibri"):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def add_section_header(doc: Document, number: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17)
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_shading(cell, "1B2D5C")
    set_cell_borders(cell, "1B2D5C", "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_styled_run(p, f"  {number}   ", bold=True, size=12, color=TEAL)
    add_styled_run(p, title.upper(), bold=True, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))
    doc.add_paragraph()


def add_question(doc: Document, q_number: str, question: str, helper: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    add_styled_run(p, f"Q{q_number}. ", bold=True, size=11, color=TEAL)
    add_styled_run(p, question, bold=True, size=11, color=NAVY)
    if helper:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(4)
        add_styled_run(hp, helper, size=9, color=GREY)
        for run in hp.runs:
            run.italic = True


def add_answer_box(doc: Document, *, height_cm: float = 1.6) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17)
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(cell, "FAFBFC")
    set_cell_borders(cell, "BFC4CC", "6")
    set_cell_height(cell, height_cm)
    p = cell.paragraphs[0]
    add_styled_run(p, " ", size=10, color=GREY)
    doc.add_paragraph()


def add_checkbox_options(doc: Document, options: list[tuple[str, str | None]]) -> None:
    for label, sub in options:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_styled_run(p, "☐  ", size=13, color=NAVY)
        add_styled_run(p, label, bold=True, size=10.5, color=NAVY)
        if sub:
            add_styled_run(p, f"  —  {sub}", size=10, color=GREY)


def add_field_row(doc: Document, label: str, height_cm: float = 0.9) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12.5)
    label_cell, input_cell = table.rows[0].cells
    label_cell.width = Cm(4.5)
    input_cell.width = Cm(12.5)
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    input_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(label_cell, "EEF2F7")
    set_cell_borders(label_cell, "BFC4CC", "6")
    set_cell_shading(input_cell, "FAFBFC")
    set_cell_borders(input_cell, "BFC4CC", "6")
    set_cell_height(label_cell, height_cm)
    lp = label_cell.paragraphs[0]
    add_styled_run(lp, f"  {label}", bold=True, size=10, color=NAVY)
    ip = input_cell.paragraphs[0]
    add_styled_run(ip, " ", size=10)


def build_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(5.5))

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_run(p, "CLIENT DISCOVERY QUESTIONNAIRE", bold=True, size=24, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_run(p, "Field Force CRM — Discovery & Requirements", size=14, color=TEAL)

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(13)
    rows = [
        ("Prepared for", "Fidson Healthcare Plc"),
        ("Prepared by", "Tales Consulting"),
        ("Document type", "Pre-engagement discovery questionnaire"),
    ]
    for i, (k, v) in enumerate(rows):
        kc, vc = table.rows[i].cells
        kc.width = Cm(4)
        vc.width = Cm(13)
        set_cell_borders(kc, "FFFFFF", "4")
        set_cell_borders(vc, "FFFFFF", "4")
        add_styled_run(kc.paragraphs[0], k, bold=True, size=10, color=GREY)
        add_styled_run(vc.paragraphs[0], v, size=10, color=NAVY)

    doc.add_paragraph()
    doc.add_paragraph()

    intro = doc.add_paragraph()
    add_styled_run(
        intro,
        "Purpose of this document",
        bold=True,
        size=12,
        color=NAVY,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_styled_run(
        p,
        "This questionnaire helps Tales Consulting understand Fidson Healthcare's existing "
        "systems, field operations, regulatory obligations, and stakeholder structure before "
        "we finalise the architecture and rollout plan for your Field Force CRM. Your answers "
        "directly inform integration design, data residency choices, access-control rules, and "
        "compliance posture.",
        size=10.5,
        color=NAVY,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_styled_run(
        p,
        "How to complete:",
        bold=True,
        size=10.5,
        color=NAVY,
    )
    add_styled_run(
        p,
        "  Type directly into the shaded answer boxes. For multi-choice questions, place an "
        "“X” inside the ☐ next to your selection. When done, save the file (or "
        "use ",
        size=10.5,
        color=NAVY,
    )
    add_styled_run(p, "File → Save As → PDF", bold=True, size=10.5, color=NAVY)
    add_styled_run(
        p,
        ") and return it to your Tales Consulting point of contact.",
        size=10.5,
        color=NAVY,
    )

    doc.add_page_break()


def build_body(doc: Document) -> None:
    # Section 1
    add_section_header(doc, "01", "Existing Systems & Integration")
    add_question(
        doc,
        "1.1",
        "Does your ERP system expose a REST API or equivalent integration interface today?",
        "If yes, please describe the API (vendor, version, documentation availability). If no, "
        "indicate whether direct database access is available for read/sync purposes.",
    )
    add_answer_box(doc, height_cm=3.2)

    add_question(
        doc,
        "1.2",
        "If no API exists — are we able to access the underlying database directly to read and sync data?",
        "Specify database type (e.g., SQL Server, Oracle, MySQL), version, and whether a read-only "
        "replica or service account can be provisioned.",
    )
    add_answer_box(doc, height_cm=2.6)

    # Section 2
    add_section_header(doc, "02", "IT Infrastructure & Hosting")
    add_question(doc, "2.1", "Who hosts and manages your current IT infrastructure?")
    add_checkbox_options(
        doc,
        [
            ("Internal IT team", "Owned and operated in-house"),
            ("Managed Service Provider (MSP)", "Outsourced to a third-party MSP"),
            ("Fully cloud-based (no in-house team)", "SaaS / cloud-native operations"),
            ("Mixed / hybrid", "Combination of the above"),
        ],
    )
    doc.add_paragraph()

    add_question(
        doc,
        "2.2",
        "Is there a preference or restriction on data residency?",
        "For example: must data be stored on Nigerian servers, on the same infrastructure as your "
        "current ERP/database, or within a specific cloud region.",
    )
    add_answer_box(doc, height_cm=2.6)

    # Section 3
    add_section_header(doc, "03", "Field Force Scale & Structure")
    add_question(
        doc,
        "3.1",
        "How many field sales representatives will use the system at launch?",
        "Please provide a number or range, plus projected growth at 12 and 36 months. This "
        "directly determines database and API architecture sizing.",
    )
    add_field_row(doc, "At launch")
    add_field_row(doc, "Projected — 12 months")
    add_field_row(doc, "Projected — 36 months")
    doc.add_paragraph()

    add_question(
        doc,
        "3.2",
        "How many territories, regions, and divisions currently exist?",
    )
    add_field_row(doc, "Territories")
    add_field_row(doc, "Regions")
    add_field_row(doc, "Divisions")
    doc.add_paragraph()

    add_question(
        doc,
        "3.3",
        "Are field reps employed directly, or are some contracted / third-party agents?",
        "This affects access control design and data ownership rules in the platform.",
    )
    add_answer_box(doc, height_cm=2.4)

    add_question(
        doc,
        "3.4",
        "Do reps work exclusively in the field, or do some have hybrid (office + field) roles?",
    )
    add_answer_box(doc, height_cm=2.4)

    # Section 4
    add_section_header(doc, "04", "Data Migration")
    add_question(
        doc,
        "4.1",
        "Is there historical sales data (orders, visits, rep performance) that needs to be imported?",
        "If yes, indicate the systems of record, approximate volume, and how far back the history "
        "extends.",
    )
    add_answer_box(doc, height_cm=3.4)

    # Section 5
    add_section_header(doc, "05", "Regulatory & Compliance")
    add_question(
        doc,
        "5.1",
        "Is the company subject to any specific regulatory framework governing field sales activity?",
        "For example: NAFDAC guidelines, internal SOPs on HCP engagement, sample distribution rules, "
        "promotional material approval requirements.",
    )
    add_answer_box(doc, height_cm=3.2)

    add_question(
        doc,
        "5.2",
        "Are there mandatory audit trail requirements for sample tracking and distribution?",
        "Sample management reconciliation is a regulatory requirement in most pharma environments. "
        "Please indicate whether NAFDAC or internal policy mandates a full chain of custody.",
    )
    add_answer_box(doc, height_cm=2.8)

    add_question(
        doc,
        "5.3",
        "Does the system need to comply with NDPR (Nigeria Data Protection Regulation)?",
        "This affects how personally identifiable data — rep GPS location, HCP data — is "
        "stored, processed, and shared.",
    )
    add_answer_box(doc, height_cm=2.4)

    # Section 6
    add_section_header(doc, "06", "Reporting & Business Intelligence")
    add_question(
        doc,
        "6.1",
        "Are there existing BI or reporting tools currently in use?",
        "If Power BI is already licensed under Microsoft 365, we can use Power BI Embedded for "
        "national dashboards at no additional licensing cost. Otherwise, we build dashboards "
        "natively — our preferred approach.",
    )
    add_answer_box(doc, height_cm=3.0)

    # Section 7
    add_section_header(doc, "07", "Stakeholders & Deployment")
    add_question(
        doc,
        "7.1",
        "Who is the primary technical point of contact and decision-maker for system design approvals on the client side?",
        "We need a named individual with authority to sign off on architecture decisions, data "
        "access, and integration credentials.",
    )
    add_field_row(doc, "Full name")
    add_field_row(doc, "Role / title")
    add_field_row(doc, "Email")
    add_field_row(doc, "Phone")
    doc.add_paragraph()

    add_question(
        doc,
        "7.2",
        "Is there an internal IT team that will be involved in deployment, or does Tales Consulting handle end-to-end infrastructure?",
    )
    add_checkbox_options(
        doc,
        [
            (
                "Tales Consulting handles end-to-end",
                "Full infrastructure setup and deployment managed by us",
            ),
            (
                "Internal IT team will be involved",
                "We coordinate with your IT team for access, firewalls, domain setup, etc.",
            ),
            ("Hybrid collaboration", "Shared responsibility — please describe below"),
        ],
    )
    add_answer_box(doc, height_cm=2.0)


def build_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(
        p,
        "Tales Consulting   •   Client Discovery Questionnaire — Fidson Healthcare   •   Confidential",
        size=8,
        color=GREY,
    )


def build_closing(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17)
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_shading(cell, "EEF7F6")
    set_cell_borders(cell, "2DCBC4", "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    add_styled_run(p, "Thank you.", bold=True, size=12, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    add_styled_run(
        p2,
        "Once completed, please save this document (or export to PDF via File → Save As → "
        "PDF) and return it to your Tales Consulting point of contact. We will use your responses "
        "to prepare a tailored architecture and engagement plan for the Fidson Field Force CRM.",
        size=10.5,
        color=NAVY,
    )


def configure_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    configure_margins(doc)
    build_cover(doc)
    build_body(doc)
    build_closing(doc)
    build_footer(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
